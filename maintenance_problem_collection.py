import tkinter as tk, os, getpass, datetime, re, math
from docx import Document
import xlwt, xlrd, xlutils.copy
from tkinter import ttk

root_path = r'\\siarnd-fs\sia01\CNP_IMSCM_F\融合产品线维护部\1_NGN联合维护组项目文件夹\100 维护组\维护一组\37 现网保障&案例写作&问题记录'
record_question = '问题录入，格式如：【陕西电信】【1137689】入局呼叫甄别失败问题'

creator_chinese = {
    'l00238065': '李广',
    'h00396308': '何军',
    'l00474792': '李梦臣',
    'l00371553': '刘坚强',
    'l00382665': '刘伟',
    'w00190780': '王勇',
    'y00478622': '殷超超',
    'z00452218': '张伟',
    'z00381447': '周伟',
}

RTAC_names = [
    '艾婧婧', '白雪天', '陈俊', '党青亮', '扈文聪', '胡夏', '李桂峰', '孙凡喜', '李槐', '刘江华', '戚小蕾', '徐苇', '徐有海', '张群'
]

questions_type_directory = [
    ('00 AGCF&SIP类问题归总', 0),
    ('01 license类问题归总', 1),
    ('02 SOSM&ETSI&GB监听类问题归总', 2),
    ('03 SSF类问题归总', 3),
    ('04 OSG问题归总', 4),
    ('05 补充业务问题归总', 5),
    ('06 彩铃传真类问题归总', 6),
    ('07 网管类问题归总', 7),
    ('08 过载类问题归总', 8),
    ('09 号码甄别&变换类问题归总', 9),
    ('10 话单类问题归总', 10),
    ('11 话统类问题归总', 11),
    ('12 扩容类问题归总', 12),
    ('13 前转类问题归总', 13),
    ('14 网关类问题归总', 14),
    ('15 消息跟踪类问题归总', 15),
    ('16 智能业务问题归总', 16),
    ('17 Uportal登录鉴权类问题', 17),
    ('18 Uportal话单类问题', 18),
    ('19 Uportal漏洞类问题', 19),
    ('20 Uportal呼叫类问题', 20),
    ('21 U-Path软终端问题归总', 21),
    ('22 PMSI酒店接口机问题归总', 22),
    ('23 CBS酒管系统问题归总', 23),
    ('24 维护资料问题', 24),
]

def get_region():
    return view_string_regin.get()

def get_product_name():
    return view_string_product.get()

def get_site_information():
    return view_string_site.get()

def get_icare_number():
    return view_string_icare.get()

def get_RTAC_name():
    return view_string_RTAC_names.get()

def get_questions_description():
    return view_string_question.get()

def get_question_detail_info():
    return text_info.get(0.0, tk.END)

def get_rfc_type():
    return view_string_rfc_type.get()

def get_rfc_product():
    return view_string_rfc_product.get()

def get_rfc_dts():
    return view_string_rfc_dts.get()

def get_rfc_detail():
    return rfc_info.get(0.0, tk.END)

def get_rfc_region():
    return view_string_rfc_region.get()

def get_rfc_site_info():
    return view_string_rfc_site.get()

def get_rfc_version():
    return view_string_rfc_version.get()

def get_case_written_description():
    return view_string_case_written.get()

def get_is_quit():
    return view_int_is_quit.get()

def get_directory_choice():
    for index in range(len(questions_type_directory) + 1):
        if view_int_choice.get() == index:
            return questions_type_directory[index][0]

def write_question_to_excel(write_type):
    global root_path
    creator = creator_chinese[getpass.getuser()]
    date_time = datetime.datetime.now().strftime('%Y-%m-%d')
    record_excel_file_name = root_path + '\\' + '现网问题记录_录入的时候会自动填写.xls'
    if 'rfc' == write_type:
        rfc_product = get_rfc_product()
        rfc_region = get_rfc_region()
        rfc_site = get_rfc_site_info()
        rfc_version = get_rfc_version()
        rfc_type = get_rfc_type()
        rfc_detail = get_rfc_detail()
        rfc_risk = '低'
        rfc_dts = get_rfc_dts()
        rfc_interface = creator
        # 打开xls格式的excel文件 #
        excel_file = xlrd.open_workbook(filename=record_excel_file_name, formatting_info=True)
        table = excel_file.sheet_by_name('RFC操作')
        # 得到当前行和列，新增数据要从nrow + 1行写入 #
        nrows = table.nrows
        ncol = table.ncols
        write_result_info = [rfc_product, rfc_region, rfc_site, rfc_version, rfc_type, rfc_detail, date_time, rfc_risk, rfc_dts, rfc_interface]
        tmp_excel_file = xlutils.copy.copy(excel_file)
        tmp_table = tmp_excel_file.get_sheet('RFC操作')
        for col in range(ncol):
            tmp_table.write(nrows, col, write_result_info[col])
        tmp_excel_file.save(record_excel_file_name)
    if 'question' == write_type:
        region = get_region()
        product_name = get_product_name()
        site_info = get_site_information()
        icare_no = get_icare_number()
        question = get_questions_description()
        question_detail = get_question_detail_info()
        if question_detail == '':
            question_detail = question
        rtac_name = get_RTAC_name()
        is_public_flag = '否'
        question_state = 'OPEN'
        # 打开xls格式的excel文件 #
        excel_file = xlrd.open_workbook(filename=record_excel_file_name, formatting_info=True)
        table = excel_file.sheet_by_name('问题录入')
        # 得到当前行和列，新增数据要从nrow + 1行写入 #
        nrows = table.nrows
        ncol = table.ncols
        write_result_info = [nrows, date_time, product_name, region, site_info, question_detail, is_public_flag, creator,
                             is_public_flag, question_state, icare_no, rtac_name, '']
        tmp_excel_file = xlutils.copy.copy(excel_file)
        tmp_table = tmp_excel_file.get_sheet('问题录入')
        for col in range(ncol):
            tmp_table.write(nrows, col, write_result_info[col])
        tmp_excel_file.save(record_excel_file_name)

def start_create_and_open():
    global root_path
    creator = getpass.getuser()
    select_path = get_directory_choice()
    site_info = get_site_information()
    case_written = get_case_written_description()
    month_str = datetime.datetime.now().strftime('%Y-%m-%d')
    # month_str = now_time.replace('-', '_')
    question = get_questions_description()
    is_quit = get_is_quit()
    rfc_region = get_rfc_region()
    if rfc_region is not '':
        write_type = 'rfc'
        write_question_to_excel(write_type)
    if case_written is not '':
        # create a description document #
        document_name = '【' + creator_chinese[creator] + '】' + case_written + '_' + month_str + '.docx'
        # 打开文档
        document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        # document = Document()
        document.add_paragraph('')
        # 保存文件 #
        case_written_path = r'\\siarnd-fs\sia01\CNP_IMSCM_F\融合产品线维护部\1_NGN联合维护组项目文件夹\100 维护组\维护一组\37 现网保障&案例写作&问题记录\共性问题案例写作'
        save_file_name = case_written_path + '\\' + document_name
        document.save(save_file_name)
        os.startfile(case_written_path)
        error_msg_path = '共性问题案例写作\\' +  document_name
        error_msg = '\n作者【%s】在【%s】时间写了一个问题案例：%s' % (creator_chinese[creator], month_str, error_msg_path)
        with open('log.txt', 'a') as log_file:
            log_file.write(error_msg)
    if site_info is not '':
        icare_number = get_icare_number()
        if site_info == '':
            site_info = 'xx'
        if icare_number == '':
            icare_number = 'xxxxx'
        # eg: 37 现网保障&案例写作&问题记录\09 号码甄别&变换类问题归总\【刘伟】【陕西电信】【SR 1137689】入局呼叫甄别失败问题_xxxxx_2019_xx_xx #
        result_path = root_path + '\\' + select_path + '\\'  + '【' + creator_chinese[creator] + '】' \
                      + '【' + site_info + '】' + '【' + icare_number + '】' + question + '_' + month_str
        print(result_path)
        # create the directory #
        os.makedirs(result_path)
        # open the directory #
        os.startfile(result_path)
        write_type = 'question'
        write_question_to_excel(write_type)
        error_path_info = select_path + '\\'  + '【' + creator_chinese[creator] + '】' \
                      + '【' + site_info + '】' + '【' + icare_number + '】' + question + '_' + month_str
        error_msg = '\n【%s】在【%s】创建了一个问题文件夹： %s' % (creator_chinese[creator], month_str, error_path_info)
        with open('log.txt', 'a') as log_file:
            log_file.write(error_msg)
    else:
        error_msg = '\n据点信息 或 问题描述不能为空， 执行失败。'
        with open('log.txt', 'a') as log_file:
            log_file.write(error_msg)
    # 默认记录一次就关闭窗口 #
    if is_quit == True:
        window.quit()


window = tk.Tk()
window.title('Problem Collection Tool')
window.resizable(False, False)
window.geometry('800x1000')

current_rows, current_column = 0, 0
view_string_regin, view_int_choice, view_string_question = tk.StringVar(), tk.IntVar(), tk.StringVar()
view_string_site, view_string_icare = tk.StringVar(), tk.StringVar()
ttk.Label(window, text='发生问题区域：').grid(row=0, padx=5, pady=3, sticky=tk.W)
view_string_regin.set('国内')
ttk.Radiobutton(window, text='国内', value='国内',  command=get_region, variable=view_string_regin).\
            grid(row=current_rows, column=current_column + 1, padx=5, sticky=tk.W)
ttk.Radiobutton(window, text='海外', value='海外', command=get_region, variable=view_string_regin).\
            grid(row=current_rows, column=current_column + 2, padx=5, sticky=tk.W)
current_rows += 1
view_string_product = tk.StringVar()
view_string_product.set('SoftX3000')
ttk.Label(window, text='产品名称：').grid(row=current_rows, padx=5, sticky=tk.W)
current_rows += 1
ttk.Radiobutton(window, text='SoftX3000', value='SoftX3000',  command=get_product_name, variable=view_string_product).\
            grid(row=current_rows, column=current_column, padx=5, sticky=tk.W)
ttk.Radiobutton(window, text='UAC3000', value='UAC3000', command=get_product_name, variable=view_string_product).\
            grid(row=current_rows, column=current_column + 1, padx=5, sticky=tk.W)
ttk.Radiobutton(window, text='Uportal', value='Uportal', command=get_product_name, variable=view_string_product).\
            grid(row=current_rows, column=current_column + 2, padx=5, sticky=tk.W)
current_rows += 1
ttk.Radiobutton(window, text='U-Path软终端', value='U-Path软终端',  command=get_product_name, variable=view_string_product).\
            grid(row=current_rows, column=current_column, padx=5, sticky=tk.W)
ttk.Radiobutton(window, text='PMSI酒店接口机', value='PMSI酒店接口机', command=get_product_name, variable=view_string_product).\
            grid(row=current_rows, column=current_column + 1, padx=5, sticky=tk.W)
ttk.Radiobutton(window, text='CBS酒管系统&OSG', value='CBS酒管系统&OSG', command=get_product_name, variable=view_string_product).\
            grid(row=current_rows, column=current_column + 2, padx=5, sticky=tk.W)

current_rows += 1
ttk.Label(window, text='选择其中一个路径添加问题', width=80).\
            grid(row=current_rows, columnspan=3, sticky=tk.W)
index = 0
current_rows += 1
for question, question_type in questions_type_directory:
    ttk.Radiobutton(window, text=question, value=question_type, command=get_directory_choice, variable=view_int_choice).\
            grid(row=current_rows, column=current_column + index, padx=5, pady=2, sticky=tk.E)
    index += 1
    if 3 == index:
        index = 0
        current_rows += 1

current_rows += 1
ttk.Label(window, text='').grid(row=current_rows, columnspan=3, sticky=tk.W)
current_rows += 1
label = ttk.Label(window,
    text=record_question,       # 标签的文字
    width=80,           # 标签长宽
    )
label.grid(row=current_rows, columnspan=3)    # 固定窗口位置
current_rows += 1
ttk.Label(window, text='局点信息：').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
ttk.Entry(window, textvariable=view_string_site, width=27).grid(row=current_rows, column=1, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='icare单号：').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
ttk.Entry(window, textvariable=view_string_icare, width=27).grid(row=current_rows, column=1, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='问题描述：').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
ttk.Entry(window, textvariable=view_string_question, width=50).grid(row=current_rows, column=1, columnspan=3, sticky=tk.W)

# 创建一个下拉列表
current_rows += 1
ttk.Label(window, text='RTAC人员：').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
view_string_RTAC_names = tk.StringVar()
numberChosen = ttk.Combobox(window, width=12, textvariable=view_string_RTAC_names)
numberChosen['values'] = tuple(RTAC_names)     # 设置下拉列表的值
numberChosen.grid(row=current_rows, column=1, sticky=tk.W)      # 设置其在界面中出现的位置
numberChosen.current(0)    # 设置下拉列表默认显示的值，0为 numberChosen['values'] 的下标值
current_rows += 1
ttk.Label(window, text='问题详细描述（选填）：').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
text_info = tk.Text(window, height=5, width=50)
# how to get Text context: text_info.get(0.0,tk.END) #
text_info.grid(row=current_rows, column=1, columnspan=2, sticky=tk.W)

current_rows += 1
ttk.Label(window, text='').grid(row=current_rows, columnspan=3, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='RFC操作--国家').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
ttk.Label(window, text='RFC操作--局点').grid(row=current_rows, column=1, padx=5, pady=2, sticky=tk.W)
ttk.Label(window, text='RFC操作--版本').grid(row=current_rows, column=2, padx=5, pady=2, sticky=tk.W)
current_rows += 1
view_string_rfc_region, view_string_rfc_site, view_string_rfc_version = tk.StringVar(), tk.StringVar(), tk.StringVar()
ttk.Entry(window, textvariable=view_string_rfc_region, width=12).grid(row=current_rows, column=0, sticky=tk.E)
ttk.Entry(window, textvariable=view_string_rfc_site, width=12).grid(row=current_rows, column=1, sticky=tk.W)
ttk.Entry(window, textvariable=view_string_rfc_version, width=30).grid(row=current_rows, column=2, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='RFC操作--具体产品').grid(row=current_rows, column=1, padx=5, pady=2, sticky=tk.W)
ttk.Label(window, text='RFC操作--操作类型').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
ttk.Label(window, text='RFC操作--变更单号').grid(row=current_rows, column=2, padx=5, pady=2, sticky=tk.W)
current_rows += 1
view_string_rfc_type, view_string_rfc_product, view_string_rfc_dts = tk.StringVar(), tk.StringVar(), tk.StringVar()
ttk.Entry(window, textvariable=view_string_rfc_type, width=12).grid(row=current_rows, column=0, sticky=tk.E)
ttk.Entry(window, textvariable=view_string_rfc_product, width=12).grid(row=current_rows, column=1, sticky=tk.W)
ttk.Entry(window, textvariable=view_string_rfc_dts, width=30).grid(row=current_rows, column=2, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='RFC操作详细描述录入：').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
rfc_info = tk.Text(window, height=7, width=60)
rfc_info.grid(row=current_rows, column=1, columnspan=2, sticky=tk.W)
#  text.insert(index,string)  index = x.y的形式,x表示行，y表示列 #
rfc_info.insert(6.0, """【操作内容】XXXXX
【业务影响】无
【实验室验证情况】无
【应急预案】操作回退
【风险描述】低
【保障措施】RFC评审，电话值守""")

current_rows += 1
ttk.Label(window, text='').grid(row=current_rows, columnspan=3, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='共性案例写作').grid(row=current_rows, column=0, padx=5, pady=2, sticky=tk.E)
view_string_case_written = tk.StringVar()
ttk.Entry(window, textvariable=view_string_case_written, width=60).grid(row=current_rows, column=1, columnspan=2, sticky=tk.W)

current_rows += 1
view_int_is_quit = tk.IntVar()
view_int_is_quit.set('1')
ttk.Radiobutton(window, text='执行一次就关闭（默认）', value='1',  command=get_is_quit, variable=view_int_is_quit).\
            grid(row=current_rows, padx=5, pady=12, sticky=tk.W)
ttk.Radiobutton(window, text='始终不关闭该软件', value='0', command=get_is_quit, variable=view_int_is_quit).\
            grid(row=current_rows, column=current_column + 1, padx=5, pady=12, sticky=tk.W)
current_rows += 1
ttk.Label(window, text='').grid(row=current_rows, columnspan=3, sticky=tk.W)
current_rows += 1
tk.Button(window, text="create and open", width='80', font=('black', 12), command=start_create_and_open, bg='green').\
    grid(row=current_rows, columnspan=3, sticky=tk.W)

window.mainloop()
