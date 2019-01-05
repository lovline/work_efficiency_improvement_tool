from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches
import tkinter as tk
import re, zipfile, os

"""
使用pyinstall生成exe之后会有个问题：docx.opc.exceptions.PackageNotFoundError: Package not found at '\docx\templates\default.docx'
可参考如下链接解决： https://www.cnblogs.com/jiayongji/p/7290410.html
"""

def zip_txt_excel_files(current_path, src_file_name, zip_dest_name):
    zip = zipfile.ZipFile(zip_dest_name, 'w', zipfile.ZIP_DEFLATED)
    sourceFileFullDir = os.path.join(current_path, src_file_name)
    zip.write(sourceFileFullDir, src_file_name)
    zip.close()

def zip_vt_result_files(vt_directory, zip_dest_name):
    sourceFiles = os.listdir(vt_directory)
    # zipFileFullDir = os.path.join(zipFilePath, fileName)
    zip = zipfile.ZipFile(zip_dest_name, 'w', zipfile.ZIP_DEFLATED)
    for sourceFile in sourceFiles:
        sourceFileFullDir = os.path.join(vt_directory, sourceFile)
        zip.write(sourceFileFullDir, sourceFile)
    zip.close()

def create_zip_files(dirPath):
    for path, dir_names, file_names in os.walk(dirPath):
        # 打包VT结果文件夹 #
        for dir_name in dir_names:
            if re.search(r'VT', dir_name):
                curr_directory = dirPath + '\\' + dir_name
                zip_dest_name = dirPath + '\\' + dir_name + '.zip'
                zip_vt_result_files(curr_directory, zip_dest_name)

        # 打包DSP_PATCH.txt和测试用例.xlsx EXCEL文件 #
        for filename in file_names:
            if re.search(r'PATCH', filename) or re.search(r'用例', filename):
                src_file_name = filename.split('.')[0]
                zip_dest_name = dirPath + '\\' + src_file_name + '.zip'
                zip_txt_excel_files(dirPath, filename, zip_dest_name)

def create_VT_document(save_path):
    # 打开文档
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 加入一级标题
    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'问题单VT 报告')
    run.font.size = Pt(24)
    run.bold  = True

    paragraph = document.add_paragraph()
    # 设置字号
    run = paragraph.add_run(u'注意如下几项：')
    run.font.size = Pt(15)
    run.bold = True
    # 增加有序列表
    document.add_paragraph(u'请使用干净的环境进行VT，防止后台有垃圾数据影响VT结果。', style='List Number')
    document.add_paragraph(u'CPCI版本重新安装BAM，打补丁进行VT，这样最保险。', style='List Number')
    document.add_paragraph(u'ATCA版本重新安装网元，打补丁进行VT，这样最保险。', style='List Number')
    document.add_paragraph(u'VT前将数据库还原成之前备份一个比较干净的数据库，再打最新的补丁。', style='List Number')
    # 增加分页
    # document.add_page_break()

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'问题单编号')
    run.font.size = Pt(15)
    run.bold = True
    # 获取路径中对应的问题单号 #
    pattern = re.compile(r'DTS\d{10,15}')
    match = pattern.search(save_path)
    if match:
        dts_number = match.group()
    else:
        dts_number = 'DTSXXXXX'
    document.add_paragraph(dts_number)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'补丁信息')
    run.font.size = Pt(15)
    run.bold = True
    # 增加无序列表
    document.add_paragraph(u'LST PATCHVER:;', style='List Bullet')
    document.add_paragraph(u'DSP PATCHVER:;', style='List Bullet')
    document.add_paragraph(u'DSP PATCH:;', style='List Bullet')
    document.add_paragraph(u'DSP INVER:;', style='List Bullet')
    document.add_paragraph(u'DSP BAMPATCH:;', style='List Bullet')
    document.add_paragraph(u'DSP SYSRES:;', style='List Bullet')
    document.add_paragraph('如下附件：')
    document.add_paragraph('')
    document.add_paragraph('')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'测试用例')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('是否完成，是否符合规范: 是，是。')
    document.add_paragraph('如下附件：')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'VT 验证结果')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('如下附件：')
    document.add_paragraph('')
    document.add_paragraph('')

    document.add_paragraph('')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'checklist项')
    run.font.size = Pt(16)
    run.bold = True
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'问题单是否走到开发组长手里（归档）？')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('是')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'热补丁资源检查')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('（注：对于块内存可能由于SIP中继心跳，或者MSGI单板自检心跳影响导致波动，可以主要关注多次测试有无异常增加）')
    document.add_paragraph('补丁前后DSP SYSRES，可参考补丁信息文件')
    document.add_paragraph('不涉及')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'热补丁异常打印检测')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('A、DeviceAlarm中搜索“call by”')
    document.add_paragraph('B、DeviceAlarm中搜索“dead”')
    document.add_paragraph('C、DeviceOSlog中搜索“destroy”')
    document.add_paragraph('不涉及')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'涉及BAM后台的修改')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('重复打一次BAM补丁能否成功')
    document.add_paragraph('不涉及')
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'检查结果是否和之前验证的结果一致?')
    run.font.size = Pt(15)
    run.bold = True
    document.add_paragraph('一致')

    # 保存文件 #
    save_file_name = save_path + '\\' + dts_number + ' VT报告' + '.docx'
    document.save(save_file_name)


def get_save_path():
    return view_string_question.get()

def start_execute():
    save_path = get_save_path()
    create_zip_files(save_path)
    create_VT_document(save_path)
    window.quit()


window = tk.Tk()
window.title('Create VT Document')
window.geometry('800x400')

view_string_question = tk.StringVar()
record_path = r'输入需要创建VT 报告的路径如 -- N:\2019问题归总\监听不触发问题'

tk.Label(window, textvariable='', width='27').pack()
tk.Label(window, textvariable='', width='27').pack()
label = tk.Label(window,
    text=record_path,       # 标签的文字
    bg='gray',                 # 背景颜色
    font=('Arial', 12),         # 字体和字体大小
    width=100, height=2          # 标签长宽
    )
label.pack()    # 固定窗口位置
tk.Label(window, textvariable='', width='27').pack()
tk.Entry(window, textvariable=view_string_question, width=100).pack()
tk.Label(window, textvariable='', width='27').pack()

tk.Label(window, textvariable='', width='27').pack()
tk.Button(window, text="open common program", height='2', width='20', font=('black', 12), command=start_execute,
           bg='#FFFAFA', fg='#4F4F4F', activebackground='white', relief='raised').pack()


window.mainloop()
